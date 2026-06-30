"""
Desktop Agent — Native OS Application Control

Interacts with the user's desktop applications (Excel, file explorer,
text editors, etc.) using PyAutoGUI for mouse/keyboard control
and direct file creation tools (openpyxl for Excel).

This agent has SMART ACTION RECIPES — it can figure out what to do
based on the task description and data from upstream agents, without
requiring pre-built action lists from the planner.
"""

import logging
import os
from pathlib import Path
from src.agents.base_agent import BaseAgent
from src.tools.desktop import DesktopTool
from src.tools.file_system import FileSystemTool
from src.tools.excel import ExcelTool
from src.config import Config
from src.utils.llm_client import LLMClient

logger = logging.getLogger(__name__)


class DesktopAgent(BaseAgent):
    """
    Controls desktop applications via native mouse/keyboard input
    and direct file creation tools.

    Capabilities:
    - Create Excel spreadsheets with data from upstream agents
    - Open applications and type text
    - Save files to the workspace
    - Keyboard shortcuts and mouse control
    """

    def execute(
        self,
        task_description: str,
        parameters: dict,
        dependency_data: dict,
    ) -> dict:
        """
        Execute a desktop automation task.

        Smart routing: if no explicit actions are provided, the agent
        analyzes the task description and parameters to determine
        what recipe to use.
        """
        self.logger.info(f"Starting desktop task: {task_description}")

        desktop_tool = DesktopTool()
        fs_tool = FileSystemTool()
        excel_tool = ExcelTool()

        try:
            # ── Smart routing: determine what to do ──────────────────
            save_to_excel = parameters.get("save_to_excel", False)
            save_to_notepad = parameters.get("save_to_notepad", False)
            save_to_word = parameters.get("save_to_word", False)
            app = parameters.get("app", "")
            actions = parameters.get("actions", [])
            content = parameters.get("content", "")

            # Auto-detect from task description or app parameter if flags aren't set
            desc_lower = task_description.lower()
            app_lower = str(app).lower() if app else ""
            if not save_to_excel and ("excel" in desc_lower or "spreadsheet" in desc_lower or app_lower == "excel"):
                save_to_excel = True
            if not save_to_notepad and ("notepad" in desc_lower or "text file" in desc_lower or app_lower == "notepad"):
                save_to_notepad = True
            if not save_to_word and ("word" in desc_lower or "docx" in desc_lower or "document" in desc_lower or "report" in desc_lower or app_lower == "word"):
                save_to_word = True

            # ── Recipe: Save to Excel ────────────────────────────────
            if save_to_excel:
                return self._excel_recipe(excel_tool, desktop_tool, dependency_data, parameters)

            # ── Recipe: Save to Word ─────────────────────────────────
            if save_to_word:
                return self._word_recipe(desktop_tool, fs_tool, dependency_data, parameters)

            # ── Recipe: Save to Notepad ──────────────────────────────
            if save_to_notepad:
                return self._notepad_recipe(desktop_tool, fs_tool, dependency_data, content, parameters)

            # ── Recipe: Open app + type content ──────────────────────
            if app and content:
                return self._open_and_type_recipe(desktop_tool, app, content)

            # ── Recipe: Open application ─────────────────────────────
            if app:
                return self._open_app_recipe(desktop_tool, app)

            # ── Recipe: Execute explicit action list ─────────────────
            if actions:
                return self._execute_actions(desktop_tool, fs_tool, actions)

            # ── Fallback: try to figure it out from dependency data ──
            if dependency_data:
                # If we have upstream data, save it to a text file
                return self._save_results_recipe(fs_tool, dependency_data, parameters)

            return {
                "success": True,
                "summary": "Desktop task completed (no specific actions needed).",
                "results": [],
            }

        except Exception as e:
            self.logger.error(f"Desktop task failed: {e}")
            return {
                "success": False,
                "summary": f"Desktop task failed: {e}",
                "results": [],
            }

    def _excel_recipe(
        self,
        excel_tool: ExcelTool,
        desktop_tool: DesktopTool,
        dependency_data: dict,
        parameters: dict,
    ) -> dict:
        """
        Create an Excel file from upstream agent data and open it.

        Extracts structured data (emails, phones, search results)
        from dependency data and creates a professional spreadsheet.
        """
        self.logger.info("Executing Excel recipe...")
        self._emit_log("📊 Creating Excel spreadsheet from research data...")

        headers = []
        rows = []

        # ── Prefer LLM-extracted structured data (from Research Agent) ───
        for dep_id, dep_result in dependency_data.items():
            if not isinstance(dep_result, dict):
                continue

            structured_data = dep_result.get("structured_data", [])
            structured_fields = dep_result.get("structured_fields", [])

            if structured_data and structured_fields:
                # Use the clean LLM-extracted data
                headers = ["#"] + [f.replace("_", " ").title() for f in structured_fields]
                if "source_url" not in structured_fields:
                    headers.append("Source URL")

                for i, item in enumerate(structured_data):
                    row = [str(i + 1)]
                    for field in structured_fields:
                        val = item.get(field, "")
                        row.append(str(val) if val is not None else "")
                    if "source_url" not in structured_fields:
                        row.append(item.get("source_url", ""))
                    rows.append(row)

                self._emit_log(
                    f"📊 Using LLM-structured data: {len(rows)} items with "
                    f"{len(structured_fields)} fields"
                )
                break  # Use the first dependency with structured data

        # ── Fallback: Extract data from upstream research tasks ──────────
        if not rows:
            for dep_id, dep_result in dependency_data.items():
                if not isinstance(dep_result, dict):
                    continue

                extracted = dep_result.get("extracted", {})
                data_items = dep_result.get("data", [])

                # Build rows from scraped sources with contact info
                emails = extracted.get("emails", [])
                phones = extracted.get("phones", [])

                if data_items:
                    # We have scraped page data
                    headers = ["#", "Source", "Title", "Snippet"]

                    # Add email and phone columns if we found any
                    if emails:
                        headers.append("Email")
                    if phones:
                        headers.append("Phone")

                    for i, item in enumerate(data_items):
                        row = [
                            str(i + 1),
                            item.get("source", item.get("url", "")),
                            item.get("title", ""),
                            (item.get("snippet", "") or item.get("content", ""))[:200],
                        ]
                        # Add email if available
                        if emails:
                            row.append(emails[i] if i < len(emails) else "")
                        if phones:
                            row.append(phones[i] if i < len(phones) else "")
                        rows.append(row)

                    # If we have more emails/phones than data items, add them
                    max_idx = len(data_items)
                    extra_count = max(len(emails), len(phones)) - max_idx
                    if extra_count > 0:
                        for j in range(extra_count):
                            idx = max_idx + j
                            row = [str(idx + 1), "", "", ""]
                            if emails:
                                row.append(emails[idx] if idx < len(emails) else "")
                            if phones:
                                row.append(phones[idx] if idx < len(phones) else "")
                            rows.append(row)

                elif emails or phones:
                    # Only contact info, no page data
                    headers = ["#"]
                    if emails:
                        headers.append("Email")
                    if phones:
                        headers.append("Phone")

                    max_len = max(len(emails), len(phones))
                    for i in range(max_len):
                        row = [str(i + 1)]
                        if emails:
                            row.append(emails[i] if i < len(emails) else "")
                        if phones:
                            row.append(phones[i] if i < len(phones) else "")
                        rows.append(row)

        # Fallback if no structured data found
        if not headers:
            headers = ["#", "Information"]
            for dep_id, dep_result in dependency_data.items():
                if isinstance(dep_result, dict):
                    summary = dep_result.get("summary", str(dep_result))
                    rows.append(["1", summary])

        if not rows:
            rows.append(["1", "No data collected from upstream tasks."])

        # ── Create the Excel file ────────────────────────────────────
        self._emit_log(f"📊 Preparing Excel file with {len(rows)} rows of data...")

        # Request approval before creating the file
        if not self.request_approval(f"Create Excel file with {len(rows)} rows of data"):
            return {
                "success": False,
                "summary": "User denied Excel file creation.",
            }

        try:
            return self._excel_recipe_visible(excel_tool, desktop_tool, dependency_data, parameters, headers, rows)
        except Exception as e:
            self.logger.error(f"Visible Excel entry failed: {e}. Falling back to background creation.")
            self._emit_log(f"⚠️ Visible Excel entry failed: {e}. Falling back to background creation.")
            return self._excel_recipe_background(excel_tool, headers, rows)

    def _excel_recipe_background(
        self,
        excel_tool: ExcelTool,
        headers: list[str],
        rows: list[list[str]],
    ) -> dict:
        """
        Background programmatic creation of Excel workbook as a fallback.
        """
        filepath = excel_tool.create_workbook(
            headers=headers,
            rows=rows,
            sheet_name="Agent Results",
        )
        excel_tool.open_in_excel(filepath)
        return {
            "success": True,
            "summary": f"Created Excel file programmatically in background: {filepath}",
            "filepath": filepath,
            "row_count": len(rows),
            "results": [{"action": "create_excel_background", "path": filepath, "rows": len(rows), "status": "done"}],
        }

    def _excel_recipe_visible(
        self,
        excel_tool: ExcelTool,
        desktop_tool: DesktopTool,
        dependency_data: dict,
        parameters: dict,
        headers: list[str],
        rows: list[list[str]],
    ) -> dict:
        """
        Create an Excel file visibly by launching Excel and typing the data cell-by-cell.
        """
        import time
        import pyautogui
        
        self.logger.info("Executing visible Excel GUI typist recipe...")
        self._emit_log("📊 Launching Excel visibly for live entry...")
        
        old_failsafe = True
        try:
            old_failsafe = pyautogui.FAILSAFE
            pyautogui.FAILSAFE = True  # Force fail-safe active for user safety
        except Exception:
            pass
            
        start_row_idx = 0
        total_rows = len(rows)
        headers_typed = False
        
        try:
            while True:
                try:
                    if not headers_typed:
                        # 1. Launch Excel
                        desktop_tool.open_application("excel")
                        time.sleep(6)  # Wait for Excel startup
                        
                        # 2. Excel standard behavior: on opening blank spreadsheet, 
                        # we need to press Enter or Esc to start a blank workbook, 
                        # or press Ctrl+N to ensure a blank sheet is created.
                        desktop_tool.press("esc")
                        time.sleep(0.5)
                        desktop_tool.hotkey("ctrl", "n")
                        time.sleep(1.5)
                        
                        # Let's focus Excel window title
                        focus_success = False
                        for title in ["Excel", "Book", "Spreadsheet", "Microsoft Excel"]:
                            if desktop_tool.press("alt"): # bypass alt key block
                                time.sleep(0.2)
                            
                            import ctypes
                            EnumWindows = ctypes.windll.user32.EnumWindows
                            EnumWindowsProc = ctypes.WINFUNCTYPE(ctypes.c_bool, ctypes.c_int, ctypes.c_int)
                            GetWindowText = ctypes.windll.user32.GetWindowTextW
                            GetWindowTextLength = ctypes.windll.user32.GetWindowTextLengthW
                            IsWindowVisible = ctypes.windll.user32.IsWindowVisible
                            SetForegroundWindow = ctypes.windll.user32.SetForegroundWindow
                            ShowWindow = ctypes.windll.user32.ShowWindow

                            found_hwnd = []
                            def foreach_window(hwnd, lParam):
                                if IsWindowVisible(hwnd):
                                    length = GetWindowTextLength(hwnd)
                                    buff = ctypes.create_unicode_buffer(length + 1)
                                    GetWindowText(hwnd, buff, length + 1)
                                    if title.lower() in buff.value.lower():
                                        found_hwnd.append(hwnd)
                                        return False
                                return True
                            
                            EnumWindows(EnumWindowsProc(foreach_window), 0)
                            if found_hwnd:
                                ShowWindow(found_hwnd[0], 9)
                                time.sleep(0.5)
                                SetForegroundWindow(found_hwnd[0])
                                time.sleep(0.5)
                                focus_success = True
                                break
                        
                        # 3. Enter A1
                        # Standard keyboard entry: Type column headers
                        self._emit_log("✍️ Typing Excel column headers live...")
                        for idx, header in enumerate(headers):
                            desktop_tool.type_text_unicode(header)
                            if idx < len(headers) - 1:
                                desktop_tool.press("tab")
                            else:
                                desktop_tool.press("enter")
                        time.sleep(0.5)
                        headers_typed = True
                    
                    # Type cells row-by-row
                    for r_idx in range(start_row_idx, total_rows):
                        row = rows[r_idx]
                        self._emit_log(f"✍️ Live typing row {r_idx + 1}/{total_rows}...")
                        for c_idx, cell in enumerate(row):
                            desktop_tool.type_text_unicode(str(cell) if cell is not None else "")
                            if c_idx < len(row) - 1:
                                desktop_tool.press("tab")
                            else:
                                desktop_tool.press("enter")
                        time.sleep(0.1) # brief pause between rows
                        start_row_idx = r_idx + 1
                        
                    # 4. Save workbook via Ctrl+S
                    self._emit_log("💾 Saving the typed Excel workbook...")
                    desktop_tool.hotkey("ctrl", "s")
                    time.sleep(3) # Wait for save dialog
                    
                    from datetime import datetime
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"agent_results_{timestamp}.xlsx"
                    filepath = str(Config.WORKSPACE_ROOT / filename)
                    
                    # Type the destination path
                    desktop_tool.type_text_unicode(filepath)
                    time.sleep(0.5)
                    desktop_tool.press("enter")
                    time.sleep(4) # Wait for save to complete
                    
                    # Press enter again in case of confirmation
                    desktop_tool.press("enter")
                    time.sleep(1)
                    
                    self._emit_log(f"✅ Excel sheet saved live: {filepath}")
                    return {
                        "success": True,
                        "summary": f"Created Excel file visibly with {len(rows)} rows: {filepath}",
                        "filepath": filepath,
                        "row_count": len(rows),
                        "results": [{"action": "create_excel_visible", "path": filepath, "rows": len(rows), "status": "done"}],
                    }
                    
                except pyautogui.FailSafeException:
                    self._emit_log(f"⚠️ Live Excel entry was halted (row index: {start_row_idx}) because the mouse was moved!")
                    
                    # First offer the continue option
                    resume = self.request_approval(
                        f"Live Excel entry was interrupted. Would you like to CONTINUE/RESUME live typing "
                        f"from row {start_row_idx + 1}? (Please select the correct cell in Excel first, then click Yes)"
                    )
                    if resume:
                        self._emit_log(f"🔄 Resuming live GUI Excel entry from row {start_row_idx + 1} in 3 seconds...")
                        time.sleep(3)
                        continue
                    
                    # Next, offer the fallback option
                    fallback_bg = self.request_approval(
                        "Would you like to fallback to creating the Excel file in the background?"
                    )
                    if fallback_bg:
                        self._emit_log("📊 Falling back to background Excel writing...")
                        return self._excel_recipe_background(excel_tool, headers, rows)
                    
                    # Next, offer to restart from the beginning
                    restart = self.request_approval(
                        "Would you like to restart the Excel visible typing from the beginning?"
                    )
                    if restart:
                        self._emit_log("🔄 Restarting live GUI Excel entry from the beginning...")
                        start_row_idx = 0
                        headers_typed = False
                        continue
                    else:
                        self._emit_log("❌ Live Excel entry aborted by user.")
                        raise RuntimeError("Excel visible entry aborted by user.")
        finally:
            try:
                pyautogui.FAILSAFE = old_failsafe
            except Exception:
                pass

    def _notepad_recipe(
        self,
        desktop_tool: DesktopTool,
        fs_tool: FileSystemTool,
        dependency_data: dict,
        content: str,
        parameters: dict,
    ) -> dict:
        """
        Save results to a text file and open in Notepad.
        """
        self.logger.info("Executing Notepad recipe...")
        self._emit_log("📝 Preparing text file from research data...")

        # Build text content from dependency data
        if not content:
            lines = ["=" * 60, "AGENT OS — RESEARCH RESULTS", "=" * 60, ""]
            for dep_id, dep_result in dependency_data.items():
                if not isinstance(dep_result, dict):
                    continue

                lines.append(f"--- Task: {dep_id} ---")
                lines.append(f"Summary: {dep_result.get('summary', 'N/A')}")
                lines.append("")

                extracted = dep_result.get("extracted", {})
                if extracted.get("emails"):
                    lines.append("Emails Found:")
                    for email in extracted["emails"]:
                        lines.append(f"  • {email}")
                    lines.append("")

                if extracted.get("phones"):
                    lines.append("Phone Numbers Found:")
                    for phone in extracted["phones"]:
                        lines.append(f"  • {phone}")
                    lines.append("")

                data = dep_result.get("data", [])
                if data:
                    lines.append("Sources:")
                    for item in data:
                        title = item.get("title", "")
                        url = item.get("source", item.get("url", ""))
                        lines.append(f"  • {title}: {url}")
                    lines.append("")

            content = "\n".join(lines)

        # Save the file
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"results_{timestamp}.txt"

        if not self.request_approval(f"Create text file: {filename}"):
            return {"success": False, "summary": "User denied text file creation."}

        filepath = fs_tool.write_file(filename, content)

        # Open in Notepad
        self._emit_log(f"📂 Opening in Notepad: {filepath}")
        desktop_tool.open_application("notepad")
        desktop_tool.wait(1)

        # Use os.startfile for reliability
        try:
            os.startfile(filepath)
        except Exception:
            pass

        return {
            "success": True,
            "summary": f"Created text file: {filepath}",
            "filepath": filepath,
            "results": [{"action": "create_text_file", "path": filepath, "status": "done"}],
        }

    def _open_and_type_recipe(self, desktop_tool: DesktopTool, app: str, content: str) -> dict:
        """Open an application and type content into it."""
        if not self.request_approval(f"Open {app} and type content"):
            return {"success": False, "summary": f"User denied opening {app}."}

        desktop_tool.open_application(app)
        desktop_tool.wait(2)
        desktop_tool.type_text_unicode(content)

        return {
            "success": True,
            "summary": f"Opened {app} and typed content.",
            "results": [{"action": "open_and_type", "app": app, "status": "done"}],
        }

    def _open_app_recipe(self, desktop_tool: DesktopTool, app: str) -> dict:
        """Open an application."""
        if not self.request_approval(f"Open application: {app}"):
            return {"success": False, "summary": f"User denied opening {app}."}

        desktop_tool.open_application(app)
        return {
            "success": True,
            "summary": f"Opened {app}.",
            "results": [{"action": "open_app", "app": app, "status": "done"}],
        }

    def _save_results_recipe(self, fs_tool: FileSystemTool, dependency_data: dict, parameters: dict) -> dict:
        """Save upstream results to a text file when no specific app is requested."""
        self.logger.info("Saving results to text file (fallback recipe)...")
        self._emit_log("💾 Saving results to file...")

        lines = ["AGENT RESULTS", "=" * 40, ""]
        for dep_id, dep_result in dependency_data.items():
            if isinstance(dep_result, dict):
                lines.append(f"Task {dep_id}: {dep_result.get('summary', 'Done')}")
        content = "\n".join(lines)

        from datetime import datetime
        filename = f"results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filepath = fs_tool.write_file(filename, content)

        return {
            "success": True,
            "summary": f"Results saved to: {filepath}",
            "filepath": filepath,
            "results": [{"action": "save_file", "path": filepath, "status": "done"}],
        }

    def _vision_verify_action(
        self,
        desktop_tool: DesktopTool,
        expected_text: str | None = None,
        timeout: float = 5.0,
    ) -> bool:
        """
        Take a screenshot after an action and verify the screen state
        changed as expected using OCR.

        Args:
            desktop_tool: The DesktopTool instance for screenshots.
            expected_text: If provided, verify this text appears on screen.
            timeout: Max seconds to wait for the expected state.

        Returns:
            True if verification passed, False otherwise.
        """
        import time
        from src.vision.ocr import OCR

        ocr = OCR()
        start = time.time()

        while time.time() - start < timeout:
            try:
                screen_text = ocr.extract_from_screen()
                if expected_text is None:
                    # No specific expectation — just confirm screen is readable
                    return len(screen_text) > 0
                if expected_text.lower() in screen_text.lower():
                    self.logger.info(
                        f"Vision verify: found expected text '{expected_text}'"
                    )
                    return True
            except Exception as e:
                self.logger.debug(f"Vision verify error: {e}")
            time.sleep(0.5)

        self.logger.warning(
            f"Vision verify: expected text '{expected_text}' not found within {timeout}s"
        )
        return False

    def _vision_find_and_click(
        self,
        desktop_tool: DesktopTool,
        target_text: str,
        timeout: float = 8.0,
    ) -> bool:
        """
        Use OCR to locate text on screen and click it.

        This is the core of the vision-guided control loop:
        1. Take a screenshot
        2. Run OCR to find the target text and its bounding box
        3. Click the center of the bounding box
        4. Verify the state changed

        Args:
            desktop_tool: DesktopTool instance.
            target_text: The text to search for on screen (e.g. a button label).
            timeout: Max seconds to search for the text.

        Returns:
            True if the text was found and clicked, False otherwise.
        """
        import time
        from src.vision.ocr import OCR

        ocr = OCR()
        start = time.time()
        self._emit_log(f"👁️ Vision: searching for '{target_text}' on screen...")

        while time.time() - start < timeout:
            try:
                coords = ocr.find_text_on_screen(target_text)
                if coords:
                    x, y = coords
                    self._emit_log(
                        f"👁️ Vision: found '{target_text}' at ({x}, {y}), clicking..."
                    )
                    desktop_tool.click(x, y)
                    time.sleep(0.5)
                    return True
            except Exception as e:
                self.logger.debug(f"Vision find error: {e}")
            time.sleep(0.5)

        self._emit_log(f"👁️ Vision: could not find '{target_text}' on screen")
        return False

    def _execute_actions(self, desktop_tool: DesktopTool, fs_tool: FileSystemTool, actions: list) -> dict:
        """Execute an explicit list of actions (legacy support)."""
        results = []

        for action in actions:
            action_type = action.get("type", "")
            self.logger.info(f"Desktop action: {action_type}")

            if action_type == "click":
                x = action.get("x", 0)
                y = action.get("y", 0)
                desktop_tool.click(x, y)
                results.append({"action": "click", "x": x, "y": y, "status": "done"})

            elif action_type == "type":
                text = action.get("text", "")
                desktop_tool.type_text_unicode(text)
                results.append({"action": "type", "status": "done"})

            elif action_type == "hotkey":
                keys = action.get("keys", [])
                desktop_tool.hotkey(*keys)
                results.append({"action": "hotkey", "keys": keys, "status": "done"})

            elif action_type == "screenshot":
                path = action.get("path", "desktop_screenshot.png")
                desktop_tool.screenshot(path)
                results.append({"action": "screenshot", "path": path, "status": "done"})

            elif action_type == "write_file":
                file_path = action.get("path", "")
                content = action.get("content", "")
                if self.request_approval(f"Write file: {file_path}"):
                    fs_tool.write_file(file_path, content)
                    results.append({"action": "write_file", "path": file_path, "status": "done"})
                else:
                    results.append({"action": "write_file", "path": file_path, "status": "denied"})

            elif action_type == "read_file":
                file_path = action.get("path", "")
                content = fs_tool.read_file(file_path)
                results.append({"action": "read_file", "path": file_path, "content": content[:2000]})

        return {
            "success": True,
            "summary": f"Executed {len(actions)} desktop actions.",
            "results": results,
        }

    def _word_recipe(
        self,
        desktop_tool: DesktopTool,
        fs_tool: FileSystemTool,
        dependency_data: dict,
        parameters: dict,
    ) -> dict:
        """
        Create a styled Word document (.docx) from upstream research data.
        """
        self.logger.info("Executing Word document recipe...")
        self._emit_log("📝 Preparing Word document from research data...")

        # ── Lazy import python-docx (install if missing) ────────────────────
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
        except ImportError:
            self._emit_log("📥 python-docx not installed. Installing automatically...")
            try:
                import sys
                import subprocess
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", "python-docx"],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
            except Exception as e:
                self.logger.error(f"Failed to install python-docx: {e}")
                return {
                    "success": False,
                    "summary": f"python-docx is not installed and auto-install failed: {e}",
                    "error": str(e),
                }

        # ── Extract structured data / fallback data for tables ───────────
        headers = []
        rows = []
        raw_goal = parameters.get("raw_goal", "Research Report")

        for dep_id, dep_result in dependency_data.items():
            if not isinstance(dep_result, dict):
                continue

            structured_data = dep_result.get("structured_data", [])
            structured_fields = dep_result.get("structured_fields", [])

            if structured_data and structured_fields:
                headers = ["#"] + [f.replace("_", " ").title() for f in structured_fields]
                for i, item in enumerate(structured_data):
                    row = [str(i + 1)]
                    for field in structured_fields:
                        val = item.get(field, "")
                        row.append(str(val) if val is not None else "")
                    rows.append(row)
                break

        # Fallback to standard sources if no structured data
        if not rows:
            for dep_id, dep_result in dependency_data.items():
                if not isinstance(dep_result, dict):
                    continue
                data_items = dep_result.get("data", [])
                if data_items:
                    headers = ["#", "Source/Url", "Title", "Snippet"]
                    for i, item in enumerate(data_items):
                        rows.append([
                            str(i + 1),
                            item.get("source", item.get("url", "")),
                            item.get("title", ""),
                            (item.get("snippet", "") or item.get("content", ""))[:150]
                        ])
                    break

        # ── LLM report outline generation ────────────────────────────────
        llm = LLMClient.get_instance()
        report_data = None
        if llm.available:
            self._emit_log("🤖 Using local LLM to draft report outline and sections...")
            prompt = (
                f"Topic: {raw_goal}\n\n"
                f"We are generating a Microsoft Word (.docx) report.\n"
                f"Write structured content for this report. Return ONLY valid JSON in this format:\n"
                f"{{\n"
                f"  \"title\": \"Report Main Title\",\n"
                f"  \"subtitle\": \"Brief report subtitle\",\n"
                f"  \"sections\": [\n"
                f"    {{\n"
                f"      \"heading\": \"Section Heading\",\n"
                f"      \"paragraphs\": [\"Paragraph 1 content...\", \"Paragraph 2...\"],\n"
                f"      \"bullets\": [\"Bullet point 1...\", \"Bullet point 2...\"]\n"
                f"    }}\n"
                f"  ]\n"
                f"}}\n"
            )
            report_data = llm.generate_json(
                prompt=prompt,
                system="You are a professional report compiler. Return ONLY JSON.",
                temperature=0.3,
                max_tokens=2048,
                timeout=60,
            )

        # Fallback report layout if LLM fails or is unavailable
        if not report_data or "sections" not in report_data:
            report_data = {
                "title": raw_goal,
                "subtitle": "Generated Report by AgentOS",
                "sections": [
                    {
                        "heading": "Executive Summary",
                        "paragraphs": [
                            "This document contains the collected findings and results retrieved "
                            "by the autonomous Multi-Agent workspace based on the user's research goal."
                        ],
                        "bullets": [
                            "Information compiled dynamically from online sources.",
                            "Processed and formatted locally to guarantee privacy."
                        ]
                    }
                ]
            }

        # ── Document styling & construction ─────────────────────────────
        doc = Document()
        
        # Apply standard style/margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Helper to apply font formatting
        def format_run(run, font_name="Segoe UI", size_pt=11, color_rgb=(51, 51, 51), bold=False, italic=False):
            run.font.name = font_name
            run.font.size = Pt(size_pt)
            run.font.color.rgb = RGBColor(*color_rgb)
            run.bold = bold
            run.italic = italic

        # Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(report_data.get("title", raw_goal))
        format_run(run_title, font_name="Segoe UI Light", size_pt=26, color_rgb=(30, 41, 59), bold=True)
        p_title.paragraph_format.space_after = Pt(4)

        # Subtitle
        subtitle_text = report_data.get("subtitle", "")
        if subtitle_text:
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sub = p_sub.add_run(subtitle_text)
            format_run(run_sub, font_name="Segoe UI", size_pt=12, color_rgb=(100, 116, 139), italic=True)
            p_sub.paragraph_format.space_after = Pt(24)

        # Divider line
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # Sections
        for sec in report_data.get("sections", []):
            heading_text = sec.get("heading", "")
            if heading_text:
                p_head = doc.add_paragraph()
                p_head.paragraph_format.space_before = Pt(18)
                p_head.paragraph_format.space_after = Pt(6)
                run_head = p_head.add_run(heading_text)
                format_run(run_head, font_name="Segoe UI Semibold", size_pt=16, color_rgb=(79, 70, 229), bold=True)

            for p_text in sec.get("paragraphs", []):
                p_para = doc.add_paragraph()
                p_para.paragraph_format.space_after = Pt(8)
                p_para.paragraph_format.line_spacing = 1.15
                run_para = p_para.add_run(p_text)
                format_run(run_para, font_name="Segoe UI", size_pt=11)

            for b_text in sec.get("bullets", []):
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(4)
                run_bullet = p_bullet.add_run(b_text)
                format_run(run_bullet, font_name="Segoe UI", size_pt=11)

        # Add Data Table
        if headers and rows:
            p_table_title = doc.add_paragraph()
            p_table_title.paragraph_format.space_before = Pt(24)
            p_table_title.paragraph_format.space_after = Pt(8)
            run_table_title = p_table_title.add_run("Collected Research Data Listing")
            format_run(run_table_title, font_name="Segoe UI Semibold", size_pt=14, color_rgb=(79, 70, 229), bold=True)

            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Light Shading Accent 1'

            # Header Row
            hdr_cells = table.rows[0].cells
            for col_idx, header in enumerate(headers):
                hdr_cells[col_idx].text = header
                # Set shading
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4F46E5"/>')
                hdr_cells[col_idx]._tc.get_or_add_tcPr().append(shading_elm)
                # Bold white font
                for p in hdr_cells[col_idx].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        format_run(run, font_name="Segoe UI Semibold", size_pt=10, color_rgb=(255, 255, 255), bold=True)

            # Data Rows
            for row_idx, r_data in enumerate(rows):
                row_cells = table.rows[row_idx + 1].cells
                for col_idx, text in enumerate(r_data):
                    row_cells[col_idx].text = str(text)
                    for p in row_cells[col_idx].paragraphs:
                        for run in p.runs:
                            format_run(run, font_name="Segoe UI", size_pt=9.5)

        # Save document
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = "".join(c for c in report_data.get("title", raw_goal) if c.isalnum() or c in (" ", "_", "-")).strip()
        safe_title = safe_title.replace(" ", "_")[:40]
        filename = f"report_{safe_title}_{timestamp}.docx"
        filepath = str(Config.WORKSPACE_ROOT / filename)

        # Ask user for approval
        if not self.request_approval(f"Save Word report as: {filename}"):
            return {
                "success": False,
                "summary": "User rejected saving MS Word report.",
                "error": "User rejected saving MS Word report",
            }

        doc.save(filepath)
        self._emit_log(f"✅ MS Word document saved successfully: {filepath}")

        # Attempt to open
        try:
            os.startfile(filepath)
            self._emit_log(f"📂 Opened Word document: {filepath}")
        except Exception as e:
            self._emit_log(f"📂 File saved at: {filepath} (Open manually)")

        return {
            "success": True,
            "summary": f"Created Word report document: {filepath}",
            "filepath": filepath,
        }
